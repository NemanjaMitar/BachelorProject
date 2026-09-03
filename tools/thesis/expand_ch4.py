# -*- coding: utf-8 -*-
"""Expand chapter 4 per mentor's feedback:
   (1) precise definitions of курикулум/штафета/биом/... at the start of ch. 4,
   (2) fully reproducible network description in 4.2: exact per-layer dims,
       architecture diagram, real code listing from PPO.py, parameter count,
       and a step-by-step update description.
   Old 'Слика 4.2' (path of level 17) becomes 'Слика 4.3'.
   Saves to a NEW file, original untouched."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as WDA

SRC = "thesis/drafts/Дипломски Jump King.docx"
OUT = "thesis/drafts/Дипломски Jump King - v2.docx"
doc = Document(SRC)

def find(pred):
    for p in doc.paragraphs:
        if pred(p.text.strip()): return p
    return None

h41    = find(lambda t: t.startswith("4.1 Модели"))
body42 = find(lambda t: t.startswith("Коришћена је сопствена имплементација PPO"))
capfig = find(lambda t: t.startswith("Слика 4.1"))
assert h41 is not None and body42 is not None
BODY = body42.style
CAP  = capfig.style if capfig else BODY

# ---- renumber old figure 4.2 -> 4.3 (body caption + front list entry) ----
for p in doc.paragraphs:
    if p.text.strip().startswith("Слика 4.2"):
        for r in p.runs:
            if "4.2" in r.text: r.text = r.text.replace("4.2", "4.3", 1); break
        else:
            if p.runs: p.runs[0].text = p.runs[0].text.replace("Слика 4.2", "Слика 4.3", 1)

# ================= (1) DEFINITIONS before 4.1 =================
def P_before(target, style=None):
    return target.insert_paragraph_before("", style=style or BODY)

intro = P_before(h41)
intro.add_run("Методологија користи неколико појмова у прецизно одређеном, техничком "
              "значењу. Да би излагање које следи било једнозначно, дефинишемо их унапред.")
DEFS = [
 ("Епизода", "једно извршавање агента од постављеног почетног стања до успеха (достигнут "
  "циљни ниво), неуспеха (пад испод почетног нивоа епизоде) или истека ограничења броја акција."),
 ("Базен почетних стања", "коначан скуп тачака (x, y) на платформама нивоа из којих се "
  "епизоде започињу. Пуни се стањима примопредаје, аутоматски генерисаним семенима и "
  "међустањима доказане путање из претраге графа (поглавље 5)."),
 ("Курикулум (програм учења)", "правило које одређује који је подскуп базена почетних стања "
  "тренутно доступан обучавању и којим се редоследом стања уводе. Обрнути курикулум је "
  "специјални случај коришћен у раду: обучавање почиње од стања најближих излазу нивоа, а "
  "сваки прелазак прага успешности откључава следеће, теже (дубље) стање — агент тако у "
  "сваком тренутку учи само један нови сегмент пута."),
 ("Приоритетно узорковање", "расподела избора почетних стања унутар откључаног подскупа: "
  "стања се бирају пропорционално експоненцијалном клизном просеку неуспеха, уз мали "
  "аддитивни под који савладаним стањима обезбеђује повремено обнављање."),
 ("Штафета", "начин извршавања целе игре у ком је сваком нивоу додељен задужени модел; "
  "посебан извршни програм у свакој тачки одлучивања пита модел задужен за тренутни ниво, "
  "а прелазак краља на други ниво аутоматски предаје управљање његовом моделу."),
 ("Примопредаја", "скуп стања у којима модел нивоа N−1 фактички доводи краља на ниво N; "
  "успешност модела N мери се управо из те расподеле стања, чиме је дефинисана и његова "
  "завршеност (поглавље 4.6)."),
 ("Биом", "низ узастопних нивоа са заједничком механиком (нпр. ветровити нивои или ледени "
  "нивои) који се, када је то оправдано, обучава једним заједничким моделом уместо моделима "
  "по појединачним нивоима."),
]
for term, desc in DEFS:
    p = P_before(h41)
    r = p.add_run(term + " — "); r.bold = True
    p.add_run(desc)

# ================= (2) 4.2: reproducible network =================
# insertion point: before the hyperparameter table (first table whose header is 'Хиперпараметар')
hyper_tbl = None
for t in doc.tables:
    try:
        if t.cell(0, 0).text.strip().startswith("Хиперпараметар"): hyper_tbl = t; break
    except Exception: pass
assert hyper_tbl is not None
anchor = hyper_tbl._tbl

def P_at(style=None, align=None):
    p = doc.add_paragraph("", style=style or BODY)
    if align is not None: p.alignment = align
    anchor.addprevious(p._p)
    return p

p = P_at(); p.add_run(
 "Ради потпуне поновљивости, архитектура је у наставку наведена слој по слој, са тачним "
 "димензијама сваког међурезултата. Опсервација стиже као један вектор who паковање: првих "
 "3·45·60 = 6075 елемената је изравнана мапа заузетости (три канала: чврсте површине, "
 "позиција краља, опасне површине), а иза њих следе скаларна обележја — у основној "
 "конфигурацији четири (x/W, y/H, индекс нивоа, укупна висина), уз два додатна за фазу ветра "
 "[sin, cos] на ветровитим и два за брзину краља [vx, vy] на леденим нивоима. Мрежа вектор "
 "интерно распакује; паковање је изабрано да би бафер, векторизована окружења и целокупан "
 "остатак петље обучавања остали идентични агенту са чисто скаларном опсервацијом.")
p = P_at(); p.add_run(
 "Конволуциони део чине три слоја Conv2d са језгром 3×3, кораком 2 и допуном 1, редом са "
 "16, 32 и 32 филтера и ReLU активацијом иза сваког. Просторне димензије се тиме преполовљавају "
 "три пута: улаз 3×45×60 → 16×23×30 → 32×12×15 → 32×6×8, што изравнато даје 1536 обележја. "
 "На њих се надовезују скаларна обележја (1536 + 4 = 1540 у основној конфигурацији), па два "
 "потпуно повезана слоја од по 128 неурона са tanh активацијом. Из последњег скривеног слоја "
 "гранају се две линеарне главе: политика (128 → |A|, softmax преко Categorical расподеле; "
 "|A| = 37 за пуни скуп акција) и вредност (128 → 1). Сви слојеви иницијализовани су "
 "ортогонално са појачањем √2 и нултим помацима, осим главе политике чије је појачање 0,01 — "
 "тиме је почетна политика приближно униформна, што чува истраживање у првим ажурирањима. "
 "Мрежа укупно има ≈233.000 параметара (од чега 197.248 у првом потпуно повезаном слоју), "
 "дели стабло између актора и критичара и иста је за све нивое; конфигурације се разликују "
 "само у броју скалара и величини излазног слоја.")
# figure
p = P_at(align=WDA.CENTER)
try: p.add_run().add_picture("figures/net_architecture.png", width=Inches(6.3))
except Exception: p.add_run("[слика недоступна: net_architecture.png]")
cp = P_at(style=CAP, align=WDA.CENTER)
r = cp.add_run("Слика 4.2: Архитектура актор-критичар мреже: конволуциони ток над мапом "
               "заузетости, конкатенација са скаларним обележјима, заједничко стабло и две главе.")
r.italic = True
# listing intro
p = P_at(); p.add_run(
 "Листинг 4.1 приказује дефиницију мреже из програмског модула PPO; у спрези са табелом 4.1 "
 "омогућава непосредну репродукцију свих резултата рада.")
CODE = """class ActorCritic(nn.Module):
    def __init__(self, obs_dim, num_actions, hidden=128,
                 grid_shape=(3, 45, 60), n_scalars=4):
        super().__init__()
        C, H, W = grid_shape
        self.grid_flat = C * H * W                    # 6075
        self.conv = nn.Sequential(
            nn.Conv2d(C, 16, 3, stride=2, padding=1), nn.ReLU(),
            nn.Conv2d(16, 32, 3, stride=2, padding=1), nn.ReLU(),
            nn.Conv2d(32, 32, 3, stride=2, padding=1), nn.ReLU(),
            nn.Flatten())                             # -> 32*6*8 = 1536
        trunk_in = 1536 + n_scalars                   # + [x, y, ниво, висина]
        self.trunk = nn.Sequential(
            nn.Linear(trunk_in, hidden), nn.Tanh(),
            nn.Linear(hidden, hidden), nn.Tanh())
        self.policy_head = nn.Linear(hidden, num_actions)
        self.value_head  = nn.Linear(hidden, 1)
        self.apply(self._init)                        # ортогонално, gain = sqrt(2)
        nn.init.orthogonal_(self.policy_head.weight, 0.01)

    def _features(self, x):                           # x: (N, 6075 + n_scalars)
        grid = x[:, :self.grid_flat].reshape(-1, *self.grid_shape)
        scal = x[:, self.grid_flat:]
        return self.trunk(torch.cat([self.conv(grid), scal], dim=1))

    def forward(self, x):
        h = self._features(x)
        return self.policy_head(h), self.value_head(h).squeeze(-1)"""
for line in CODE.split("\n"):
    cpp = P_at()
    run = cpp.add_run(line if line else " ")
    run.font.name = "Consolas"; run.font.size = Pt(8)
    pf = cpp.paragraph_format
    pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = 1.0
lcap = P_at(style=CAP, align=WDA.CENTER)
r = lcap.add_run("Листинг 4.1: Дефиниција актор-критичар мреже (модул PPO, PyTorch).")
r.italic = True
# update-loop paragraph
p = P_at(); p.add_run(
 "Једно ажурирање тече на следећи начин. Осам паралелних окружења изврши по 256 макро-акција "
 "тренутном политиком; сваки прелаз (опсервација, акција, награда, стари логаритам вероватноће, "
 "процена вредности) уписује се у бафер. По завршетку закупа предности се рачунају GAE поступком "
 "једним пролазом уназад, уз бутстраповање вредности на превременим прекидима а нулту вредност "
 "на стварним терминацијама. Скуп од 2048 прелаза се затим измеша и подели у четири мини-серије, "
 "па се кроз четири епохе минимизује збир одсеченог сурогат-губитка политике, губитка вредности "
 "(тежина 0,5, такође са одсецањем) и ентропијског бонуса (тежина 0,06), уз глобално одсецање "
 "норме градијента на 0,5 и Adam оптимизатор. Хиперпараметри су сумирани у табели 4.1.")

doc.save(OUT)
print("SAVED", OUT)
