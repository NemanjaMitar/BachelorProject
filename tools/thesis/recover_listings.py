# -*- coding: utf-8 -*-
"""Recover a clean thesis: start from healthy v4, re-insert the two code listings
(reward -> 3.5, GAE -> 4.3) that were lost when v5/v6 got field-corrupted.
Pure python-docx (produces real w:t text, no Word field update). Saves v7."""
import copy
from docx import Document
from docx.shared import Pt

SRC = "thesis/drafts/Дипломски Jump King - v4.docx"
OUT = "thesis/drafts/Дипломски Jump King - v7 (recovered).docx"

CODE_FONT = "Consolas"
CODE_SIZE = Pt(8)

REWARD = r'''def step(self, action_idx):
    self._apply_action(int(action_idx))
    level = self.levels.current_level
    alt   = self._altitude()

    reward = 0.0
    d_level = level - self._prev_level
    if d_level > 0:
        reward += self.level_reward * d_level
    elif d_level < 0:
        reward += self.level_penalty * d_level
    self._prev_level = level

    if alt > self._best_alt:
        reward += (alt - self._best_alt) * self.altitude_breadcrumb
        self._best_alt = alt

    terminated = False
    if self.levels.ending:
        reward += 100.0; terminated = True
    elif self.goal_level is not None and level >= self.goal_level:
        reward += self.level_reward; terminated = True'''

REWARD_CAP = ("Листинг 3.1: Функција награде (модул JK_Env). Прелази нивоа доминирају, "
              "мрвица висине даје градијент унутар нивоа, а понављање и пад дају нулу.")

GAE = r'''def compute_gae(self, last_value, gamma=0.99, lam=0.95):
    adv = torch.zeros(self.T, self.N)
    last_gae = torch.zeros(self.N)
    for t in reversed(range(self.T)):
        next_value = last_value if t == self.T-1 else self.values[t+1]
        done = self.terminated[t]
        nonterminal = 1.0 - done
        cut = torch.clamp(self.terminated[t] + self.truncated[t], max=1.0)
        delta = self.rewards[t] + gamma*next_value*nonterminal - self.values[t]
        last_gae = delta + gamma*lam*(1.0 - cut)*last_gae
        adv[t] = last_gae
    returns = adv + self.values
    return adv, returns'''

GAE_CAP = ("Листинг 4.2: GAE процена предности (модул PPO). Строго разликује стварну "
           "терминацију епизоде од прекида због временског ограничења.")

doc = Document(SRC)


def find(pred):
    for p in doc.paragraphs:
        if pred(p.text.strip()):
            return p
    return None


# capture the exact caption formatting used by Листинг 4.1
cap41 = find(lambda t: t.startswith("Листинг 4.1:"))
cap_run = cap41.runs[0] if cap41 and cap41.runs else None
cap_italic = bool(cap_run.italic) if cap_run else False
cap_bold = bool(cap_run.bold) if cap_run else False
cap_size = cap_run.font.size if (cap_run and cap_run.font.size) else None
cap_name = cap_run.font.name if cap_run else None


def code_para_after(prev):
    """insert a blank Normal paragraph right after prev, return it"""
    new = copy.deepcopy(prev._p)
    # strip content
    for child in list(new):
        if child.tag.endswith('}pPr'):
            continue
        new.remove(child)
    prev._p.addnext(new)
    from docx.text.paragraph import Paragraph
    return Paragraph(new, prev._parent)


def add_code_line(prev, text):
    p = code_para_after(prev)
    p.style = doc.styles["Normal"]
    r = p.add_run(text if text else " ")
    r.font.name = CODE_FONT
    r.font.size = CODE_SIZE
    return p


def add_caption(prev, text):
    p = code_para_after(prev)
    p.style = doc.styles["Normal"]
    r = p.add_run(text)
    r.italic = cap_italic
    r.bold = cap_bold
    if cap_size:
        r.font.size = cap_size
    if cap_name:
        r.font.name = cap_name
    return p


def insert_listing(anchor_para, code, caption):
    """insert code lines + caption immediately after anchor_para, in order"""
    prev = anchor_para
    for line in code.split("\n"):
        prev = add_code_line(prev, line.rstrip())
    prev = add_caption(prev, caption)
    return prev


# --- reward listing: after the reward-description paragraph in 3.5 ---
reward_desc = find(lambda t: t.startswith("Функција награде намерно"))
assert reward_desc is not None, "reward desc not found"
insert_listing(reward_desc, REWARD, REWARD_CAP)

# --- GAE listing: after 'Једно ажурирање тече' paragraph in 4.3 ---
gae_desc = find(lambda t: t.startswith("Једно ажурирање тече"))
assert gae_desc is not None, "gae desc not found"
insert_listing(gae_desc, GAE, GAE_CAP)

doc.save(OUT)
print("SAVED", OUT)
